import os
from typing import Dict, Optional
import PyPDF2
import docx
import openpyxl
import markdown

class DocumentParser:
    """文档解析器"""
    
    @staticmethod
    def parse_file(file_path: str, filename: str) -> Optional[Dict[str, str]]:
        """
        解析文件内容
        
        Args:
            file_path: 文件路径
            filename: 文件名
            
        Returns:
            {'title': str, 'content': str} 或 None
        """
        ext = filename.rsplit('.', 1)[-1].lower()
        
        try:
            if ext == 'pdf':
                return DocumentParser._parse_pdf(file_path, filename)
            elif ext == 'docx':
                return DocumentParser._parse_docx(file_path, filename)
            elif ext == 'txt':
                return DocumentParser._parse_txt(file_path, filename)
            elif ext in ['xlsx', 'xls']:
                return DocumentParser._parse_excel(file_path, filename)
            elif ext == 'md':
                return DocumentParser._parse_markdown(file_path, filename)
            else:
                return None
        except Exception as e:
            print(f"Error parsing file {filename}: {str(e)}")
            return None
    
    @staticmethod
    def _parse_pdf(file_path: str, filename: str) -> Dict[str, str]:
        """解析PDF文件"""
        content = []
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text:
                    content.append(text)
        
        title = filename.rsplit('.', 1)[0]
        return {
            'title': title,
            'content': '\n\n'.join(content)
        }
    
    @staticmethod
    def _parse_docx(file_path: str, filename: str) -> Dict[str, str]:
        """解析DOCX文件"""
        doc = docx.Document(file_path)
        content = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                content.append(para.text)
        
        # 尝试从表格中提取内容
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    content.append(row_text)
        
        title = filename.rsplit('.', 1)[0]
        return {
            'title': title,
            'content': '\n\n'.join(content)
        }
    
    @staticmethod
    def _parse_txt(file_path: str, filename: str) -> Dict[str, str]:
        """解析TXT文件"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        title = filename.rsplit('.', 1)[0]
        return {
            'title': title,
            'content': content
        }
    
    @staticmethod
    def _parse_excel(file_path: str, filename: str) -> Dict[str, str]:
        """解析Excel文件"""
        workbook = openpyxl.load_workbook(file_path, data_only=True)
        content = []
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            content.append(f"=== {sheet_name} ===")
            
            for row in sheet.iter_rows(values_only=True):
                row_text = ' | '.join(str(cell) if cell is not None else '' for cell in row)
                if row_text.strip(' |'):
                    content.append(row_text)
        
        title = filename.rsplit('.', 1)[0]
        return {
            'title': title,
            'content': '\n'.join(content)
        }
    
    @staticmethod
    def _parse_markdown(file_path: str, filename: str) -> Dict[str, str]:
        """解析Markdown文件"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            md_content = f.read()
        
        # 转换为纯文本
        html = markdown.markdown(md_content)
        # 简单去除HTML标签
        import re
        text = re.sub('<[^<]+?>', '', html)
        
        title = filename.rsplit('.', 1)[0]
        return {
            'title': title,
            'content': text
        }


